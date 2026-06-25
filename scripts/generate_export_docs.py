#!/usr/bin/env python3
"""Generate Word and PDF export documents for Zstate benchmark framework."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

ROOT = Path(__file__).resolve().parents[1]
EXPORT = ROOT / "docs" / "export"
DEFINITIONS = ROOT / "docs" / "EQUITY_RESEARCH_TASK_DEFINITIONS.md"


def parse_task_definitions() -> list[dict]:
    text = DEFINITIONS.read_text(encoding="utf-8")
    tasks: list[dict] = []
    blocks = re.split(r"\n(?=### [A-Z0-9-]+:)", text)
    for block in blocks:
        m = re.match(r"### ([A-Z0-9-]+):\s*(.+)", block)
        if not m:
            continue
        task_id, name = m.group(1), m.group(2).strip()
        if not re.match(r"^[A-Z]", task_id.split("-")[0]):
            continue

        def field(label: str) -> str | None:
            fm = re.search(rf"\*\*{re.escape(label)}\*\*\s*\|\s*(.+)", block)
            return fm.group(1).strip() if fm else None

        desc_m = re.search(r"\*\*Description\*\*\s*\|\s*(.+)", block)
        if not desc_m:
            desc_m = re.search(r"\*\*Description\*\*\s*\|\s*(.+?)(?:\n\| \*\*|$)", block, re.S)
        description = desc_m.group(1).strip() if desc_m else name

        tasks.append({
            "id": task_id,
            "name": name,
            "description": description.replace("\n", " "),
            "inputs": field("Inputs") or field("Input") or "—",
            "outputs": field("Outputs") or field("Output") or "—",
            "dependencies": field("Dependencies") or "—",
            "pass": field("Pass") or "—",
            "fail": field("Fail") or "—",
            "difficulty": field("Difficulty") or "—",
        })
    return tasks


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_heading(title, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared for: Zstate.ai | Confidential")
    doc.add_paragraph("")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph("")


def build_complete_framework(tasks: list[dict]) -> Document:
    doc = Document()
    add_title(
        doc,
        "Zstate Equity Research Agent Benchmark",
        "Complete Framework — 4 Phases & 185 Task Definitions (v0.2)",
    )

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Zstate builds agentic AI datasets — Task, Trajectory, and Reward loops — using "
        "credentialed domain experts. This document defines the complete equity research "
        "benchmark: four implementation phases, three-layer scoring, 185 micro-tasks from "
        "SEC/IR data through financial models to compliant investment output, and the "
        "15-task MVD pilot scope."
    )

    add_heading(doc, "2. Four-Phase Framework", 1)

    phases = [
        (
            "Phase 1 — Foundation & Data Corpus",
            "Weeks 1–2",
            [
                "Map tickers to CIK; ingest SEC EDGAR (10-K, 10-Q, 8-K)",
                "Ingest earnings transcripts (API primary, IR fallback)",
                "Load FX reference rates (FRED); build section index",
                "Expert spot-check 10% of transcripts; lock corpus manifest",
            ],
            "Deliverable: corpus_pilot_manifest.json (~40–60 docs for 5 companies)",
        ),
        (
            "Phase 2 — Task Construction & Expert Validation",
            "Weeks 2–8",
            [
                "Author tasks from archetype templates (footnote, guidance, FX)",
                "Build ground truth with 100% citation coverage",
                "Define gold minimal section sets (not brittle tool sequences)",
                "Write Layer 1 Python verification scripts",
                "CFA peer review and publish to Task Registry",
            ],
            "Deliverable: 15 MVD tasks (5 companies × 3 archetypes); scale to 185 over v0.1–v0.5",
        ),
        (
            "Phase 3 — Agent Evaluation & Trajectory Capture",
            "Weeks 3–10",
            [
                "Run model-agnostic agents via adapter layer",
                "Execute task-type workflows (F/M/C — not universal 4-stage on all tasks)",
                "Log full trajectories: thought, tool I/O, sections accessed",
                "Detect fractures: LOOP_TOOL, SIGN_ERR, HALLUC_FILL, etc.",
                "3 runs per task per model for pilot calibration",
            ],
            "Deliverable: trajectory JSONL per run; fracture report",
        ),
        (
            "Phase 4 — Scoring, Release & Trajectory Dataset",
            "Weeks 8–10+",
            [
                "Layer 1: programmatic numeric verification",
                "Layer 2: automated section recall + expert calibration sample",
                "Layer 3: citation audit; FINRA/mandates on Type C tasks only",
                "Publish benchmark_v0.1 with leaderboard",
                "Export curated trajectories for Phase 2 training product",
            ],
            "Deliverable: reward vectors, leaderboard, trajectory dataset for SFT/RL",
        ),
    ]

    for name, timeline, steps, deliverable in phases:
        add_heading(doc, name, 2)
        doc.add_paragraph(f"Timeline: {timeline}")
        add_bullets(doc, steps)
        doc.add_paragraph(f"Deliverable: {deliverable}")

    add_heading(doc, "3. Task-Type Taxonomy (Agent Workflow)", 1)
    add_table(
        doc,
        ["Type", "Stages", "Output", "Compliance", "Example"],
        [
            ["F — Forensics", "1–2: Ingest + Extract/Verify", "Reconciliation report", "Citations only", "GOOGL footnote Q1 2026"],
            ["M — Modeling", "1–3: Ingest + Python model", "Model + assumption log", "Optional mandate", "PEP FX organic growth"],
            ["C — Coverage", "1–4: Full workflow + memo", "Investment memo + reco", "FINRA + mandate", "Full initiation (v0.5)"],
        ],
    )

    add_heading(doc, "4. Three-Layer Reward Model", 1)
    add_table(
        doc,
        ["Layer", "Weight (Type F)", "Method", "Examples"],
        [
            ["L1 — Hard accuracy", "55%", "Programmatic", "Exact values, sign checks, Python verify"],
            ["L2 — Expert judgment", "25%", "Section recall + rules", "Footnote accessed, assumption bounds"],
            ["L3 — Trust", "20%", "Citation audit", "doc_id + page + snippet; no hallucinated fill"],
        ],
    )

    add_heading(doc, "5. MVD Pilot — 15 Tasks", 1)
    mvd_rows = [
        ["GOOGL", "Footnote", "F", "Q1 2026 segment vs consolidated (hedging loss)"],
        ["GOOGL", "Guidance drift", "F", "Call guidance vs 10-Q actuals"],
        ["GOOGL", "FX organic growth", "M", "Constant-currency geographic growth"],
        ["AMZN", "Footnote", "F", "AWS segment / SBC footnote"],
        ["AMZN", "Guidance drift", "F", "Fulfillment cost commentary vs actuals"],
        ["AMZN", "FX organic growth", "M", "International segment FX-adjusted growth"],
        ["NFLX", "Footnote", "F", "Content amort vs cash spend"],
        ["NFLX", "Guidance drift", "F", "Subscriber/content guidance vs actuals"],
        ["NFLX", "FX organic growth", "M", "Regional revenue CC growth"],
        ["PEP", "Footnote", "F", "Segment vs distribution footnote"],
        ["PEP", "Guidance drift", "F", "Commodity cost outlook vs margins"],
        ["PEP", "FX organic growth", "M", "EMEA/AMESA organic growth"],
        ["KO", "Footnote", "F", "Segment reconciliation"],
        ["KO", "Guidance drift", "F", "Volume/pricing guidance vs actuals"],
        ["KO", "FX organic growth", "M", "International CC growth"],
    ]
    add_table(doc, ["Company", "Archetype", "Type", "Focus"], mvd_rows)

    add_heading(doc, "6. Full Task Catalog — 185 Tasks (Detailed)", 1)
    doc.add_paragraph(
        "Each task below includes ID, name, description, inputs, outputs, dependencies, "
        "pass/fail criteria, and difficulty. Organized by layer."
    )

    layer_names = {
        "D": "Layer 0 — Data Acquisition",
        "FS": "Layer 1 — Financial Statement Intelligence",
        "M3": "Layer 2A — Three-Statement Model",
        "DCF": "Layer 2B — DCF Valuation",
        "COMP": "Layer 2C — Trading Comps",
        "LBO": "Layer 2D — LBO Model",
        "SOTP": "Layer 2E — Sum-of-the-Parts",
        "DDM": "Layer 2F — Dividend Discount Model",
        "MA": "Layer 2G — M&A Accretion/Dilution",
        "VAL": "Layer 3 — Valuation Synthesis",
        "IND": "Layer 4 — Industry Analysis",
        "EARN": "Layer 5 — Earnings Workflow",
        "TH": "Layer 6 — Investment Thesis",
        "RISK": "Layer 7 — Risk & Special Situations",
        "CPL": "Layer 8 — Compliance",
        "OUT": "Layer 9 — Output Assembly",
    }

    current_layer = None
    for t in tasks:
        prefix = t["id"].split("-")[0]
        if prefix != current_layer:
            current_layer = prefix
            add_heading(doc, layer_names.get(prefix, prefix), 2)

        add_heading(doc, f"{t['id']}: {t['name']}", 3)
        doc.add_paragraph(f"Description: {t['description']}")
        doc.add_paragraph(f"Inputs: {t['inputs']}")
        doc.add_paragraph(f"Outputs: {t['outputs']}")
        doc.add_paragraph(f"Dependencies: {t['dependencies']}")
        doc.add_paragraph(f"Pass criteria: {t['pass']}")
        doc.add_paragraph(f"Fail conditions: {t['fail']}")
        doc.add_paragraph(f"Difficulty: {t['difficulty']}")

    add_heading(doc, "7. Roadmap to Full Catalog", 1)
    add_table(
        doc,
        ["Release", "Tasks", "Adds"],
        [
            ["v0.1 MVD", "15", "Type F + M pilot (5 companies)"],
            ["v0.1b", "45", "Scale to 15 companies"],
            ["v0.2", "+15", "3-statement model bundles"],
            ["v0.3", "+15", "DCF + comps (+ market data)"],
            ["v0.4", "+30", "LBO, SOTP, DDM"],
            ["v0.5", "+20", "Type C full initiation"],
        ],
    )

    return doc


def build_senior_process_mapping() -> Document:
    doc = Document()
    add_title(
        doc,
        "Zstate Equity Research Agent Benchmark",
        "Process Mapping — Senior Team Brief (v0.2)",
    )

    add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "This document provides process mapping, data flows, roles, and timeline for senior "
        "stakeholders. It does not include the full 185-task catalog — that is maintained "
        "separately for domain experts and task engineers."
    )

    add_heading(doc, "2. Strategic Context", 1)
    add_bullets(
        doc,
        [
            "Zstate product: Task + Trajectory + Reward datasets for agentic AI training",
            "Benchmark (Phase 1) proves task quality; trajectories (Phase 2) are the product",
            "Differentiation: credentialed experts + real SEC filings + auditable agent runs",
            "MVD: 15 eval tasks, 5 companies, 10-week pilot — not 45 tasks on day one",
        ],
    )

    add_heading(doc, "3. Four-Phase Process Map", 1)

    add_heading(doc, "Phase 1 — Foundation & Data Corpus", 2)
    doc.add_paragraph("Objective: Single source of truth for filings, transcripts, FX.")
    add_table(
        doc,
        ["Step", "Owner", "Input", "Output"],
        [
            ["1.1 CIK mapping", "Data Ops", "5 tickers", "CIK registry"],
            ["1.2 EDGAR ingest", "Data Ops", "10-K/10-Q", "Raw + indexed corpus"],
            ["1.3 Transcript ingest", "Data Ops", "Transcript API + IR fallback", "Transcript registry"],
            ["1.4 Section index", "Platform Eng", "Raw filings", "Searchable sections/tables"],
            ["1.5 Corpus QA", "Lead CFA", "10% sample", "Signed corpus manifest"],
        ],
    )

    add_heading(doc, "Phase 2 — Task Construction", 2)
    doc.add_paragraph("Objective: Expert-authored tasks with ground truth and gold paths.")
    add_table(
        doc,
        ["Step", "Owner", "Input", "Output"],
        [
            ["2.1 Template design", "Lead CFA", "3 archetypes", "Task templates"],
            ["2.2 Pilot task (GOOGL)", "Associate", "Corpus", "1 validated task"],
            ["2.3 Batch authoring", "Associate", "Templates", "15 task specs"],
            ["2.4 L1 scripts", "Associate + Eng", "Ground truth", "Verification scripts"],
            ["2.5 Peer review", "Lead CFA", "Draft tasks", "Published tasks"],
        ],
    )

    add_heading(doc, "Phase 3 — Agent Evaluation", 2)
    doc.add_paragraph("Objective: Model-agnostic eval with full trajectory capture.")
    add_table(
        doc,
        ["Step", "Owner", "Input", "Output"],
        [
            ["3.1 Adapter setup", "Platform Eng", "Model APIs", "Model registry"],
            ["3.2 Tool sandbox", "Platform Eng", "Corpus", "Search, Parser, Python"],
            ["3.3 Run campaign", "Platform Eng", "15 tasks × N models × 3", "Trajectories"],
            ["3.4 Fracture analysis", "Product + Eng", "Trajectories", "Fracture report"],
        ],
    )

    add_heading(doc, "Phase 4 — Scoring & Release", 2)
    doc.add_paragraph("Objective: Scored benchmark + trajectory export for training.")
    add_table(
        doc,
        ["Step", "Owner", "Input", "Output"],
        [
            ["4.1 Layer 1 scoring", "Platform Eng", "Agent output + GT", "Hard scores"],
            ["4.2 Layer 2 scoring", "Expert (sample)", "Trajectories", "Judgment scores"],
            ["4.3 Layer 3 audit", "Platform Eng", "Citations", "Trust scores"],
            ["4.4 Release", "Product", "All artifacts", "benchmark_v0.1 + leaderboard"],
            ["4.5 Trajectory export", "Product", "Ranked runs", "Training dataset v1"],
        ],
    )

    add_heading(doc, "4. End-to-End Data Flow", 1)
    flow = (
        "SEC EDGAR + Transcripts + FX  →  Corpus Service  →  Task Registry (prompt + GT)\n"
        "  →  Agent (via adapter)  →  Trajectory Logger  →  Scoring Engine (L1/L2/L3)\n"
        "  →  Benchmark Record  →  Leaderboard + Trajectory Dataset"
    )
    doc.add_paragraph(flow)

    add_heading(doc, "5. Task Summary (Not Full Catalog)", 1)
    add_table(
        doc,
        ["Layer", "Count", "Purpose"],
        [
            ["Data acquisition", "24", "SEC, IR, peers, FX, citations"],
            ["Financial statements", "24", "Normalization, footnotes, reconciliation"],
            ["Financial models", "79", "3-stmt, DCF, comps, LBO, SOTP, DDM, M&A"],
            ["Valuation + industry + earnings", "27", "Football field, Porter, guidance drift"],
            ["Thesis + risk + compliance + output", "31", "Reco, risks, FINRA, initiation memo"],
            ["Total catalog", "185", "Full analyst workflow indexed"],
            ["MVD pilot (now)", "15", "5 cos × 3 archetypes — Type F and M only"],
        ],
    )

    add_heading(doc, "6. Roles & Capacity", 1)
    add_table(
        doc,
        ["Role", "Hours/week", "Responsibility"],
        [
            ["Lead CFA", "10–12", "Standards, review, publish approval"],
            ["MBA Associate", "12–15", "Task authoring, ground truth, citations"],
            ["Compliance (shared)", "3–5", "FINRA/mandate rules (Type C tasks)"],
            ["Platform Engineer", "Contract/part-time", "Corpus, eval runner, L1 scoring"],
        ],
    )

    add_heading(doc, "7. Timeline (10 Weeks — MVD)", 1)
    add_table(
        doc,
        ["Week", "Phase", "Milestone"],
        [
            ["1–2", "1 + 2", "Corpus for 5 cos; GOOGL pilot task complete"],
            ["3", "2", "Templates from pilot; authoring begins"],
            ["4–6", "2 + 3", "15 tasks authored; eval runner live"],
            ["7–8", "2 + 4", "Expert review; corpus lock"],
            ["9–10", "3 + 4", "Eval campaign; fracture report; release"],
        ],
    )

    add_heading(doc, "8. Success Metrics (MVD)", 1)
    add_bullets(
        doc,
        [
            "15 published tasks with 100% ground truth citations",
            "≥2 models evaluated with full trajectory capture",
            "≥15pt score spread between best and worst model",
            "≥8 distinct fracture codes documented",
            "Expert κ ≥ 0.7 on 5-task calibration set",
        ],
    )

    add_heading(doc, "9. Decisions Required from Senior Team", 1)
    add_bullets(
        doc,
        [
            "Approve 15-task MVD scope (vs 45 on day one)",
            "Confirm expert resourcing (CFA lead + associate)",
            "Approve transcript API budget (~$100–150/mo for pilot)",
            "Confirm benchmark → trajectory dataset as product path",
            "Sign off on 10-week pilot timeline",
        ],
    )

    return doc


class PDFReport(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def sanitize_for_pdf(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2192": "->",
        "\u2265": ">=",
        "\u2264": "<=",
        "\u2212": "-",
        "\u00d7": "x",
        "\u00b1": "+/-",
        "\u03ba": "k",
        "\u2022": "-",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("ascii", "replace").decode("ascii")


def pdf_multiline(pdf: FPDF, height: int, text: str) -> None:
    """Write wrapped text with reliable left-margin reset."""
    if not text:
        return
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(
        pdf.epw,
        height,
        text,
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
    )


def docx_to_pdf_via_fpdf(doc_path: Path, pdf_path: Path, title: str) -> None:
    """Build PDF from same content structure — simplified text export."""
    from docx import Document as DocxDocument

    d = DocxDocument(doc_path)
    pdf = PDFReport()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf_multiline(pdf, 10, sanitize_for_pdf(title))
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)

    for para in d.paragraphs:
        text = sanitize_for_pdf(para.text.strip())
        if not text:
            pdf.ln(3)
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style:
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf_multiline(pdf, 8, text)
            pdf.set_font("Helvetica", "", 10)
        elif "Heading 2" in style:
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf_multiline(pdf, 7, text)
            pdf.set_font("Helvetica", "", 10)
        elif "Heading 3" in style:
            pdf.set_font("Helvetica", "B", 11)
            pdf_multiline(pdf, 6, text)
            pdf.set_font("Helvetica", "", 10)
        else:
            pdf_multiline(pdf, 5, text)

    for table in d.tables:
        pdf.ln(2)
        for row in table.rows:
            line = sanitize_for_pdf(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
            if len(line) > 500:
                line = line[:497] + "..."
            pdf.set_font("Helvetica", "", 8)
            pdf_multiline(pdf, 4, line)
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(2)

    pdf.output(str(pdf_path))


def main() -> None:
    EXPORT.mkdir(parents=True, exist_ok=True)
    tasks = parse_task_definitions()
    print(f"Parsed {len(tasks)} tasks")

    # Document 1 — Complete framework
    doc1 = build_complete_framework(tasks)
    docx1 = EXPORT / "ZSTATE_Complete_Framework_185_Tasks.docx"
    pdf1 = EXPORT / "ZSTATE_Complete_Framework_185_Tasks.pdf"
    doc1.save(docx1)
    docx_to_pdf_via_fpdf(docx1, pdf1, "Zstate Complete Framework — 4 Phases & 185 Tasks")
    print(f"Wrote {docx1}")
    print(f"Wrote {pdf1}")

    # Document 2 — Senior process mapping
    doc2 = build_senior_process_mapping()
    docx2 = EXPORT / "ZSTATE_Process_Mapping_Senior_Team.docx"
    pdf2 = EXPORT / "ZSTATE_Process_Mapping_Senior_Team.pdf"
    doc2.save(docx2)
    docx_to_pdf_via_fpdf(docx2, pdf2, "Zstate Process Mapping — Senior Team Brief")
    print(f"Wrote {docx2}")
    print(f"Wrote {pdf2}")


if __name__ == "__main__":
    main()
