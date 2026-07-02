#!/usr/bin/env python3
"""Generate team-share pilot status brief (Word + PDF) for Zstate WIP."""

from __future__ import annotations

import json
import subprocess
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from statistics import median

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

ROOT = Path(__file__).resolve().parents[1]
BENCH = ROOT / "benchmark_v0.1"
EXPORT = ROOT / "docs" / "export"
PILOT_CAMPAIGN_ID = "pilot_eval_5task_v1"

TASK_SHORT_NAMES = {
    "PEP_fx_organic_growth": "PEP FX / organic growth",
    "AMZN_footnote_reconciliation": "AMZN segment bridge",
    "NFLX_guidance_drift": "NFLX guidance drift",
    "KO_footnote_reconciliation": "KO segment bridge",
    "GOOGL_footnote_reconciliation": "GOOGL segment sum (ceiling)",
}


def git_head() -> str:
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=ROOT,
            text=True,
        ).strip()
    except (subprocess.CalledProcessError, FileNotFoundError):
        return "unknown"


def load_json(path: Path) -> dict | list | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def task_summaries() -> list[dict]:
    manifest = load_json(BENCH / "manifest.json") or {}
    rows: list[dict] = []
    for entry in manifest.get("pilot_tasks", []):
        if entry.get("status") != "published":
            continue
        task_path = BENCH / entry["paths"]["task"]
        task = load_json(task_path) or {}
        prompt = (task.get("prompt") or {}).get("text", "")
        first_line = prompt.split("\n", 1)[0].strip() if prompt else entry["task_id"]
        rows.append({
            "task_id": entry["task_id"],
            "ticker": task.get("ticker", "—"),
            "type": entry.get("task_type", task.get("task_type", "—")),
            "tier": task.get("difficulty_tier", "—"),
            "summary": first_line,
        })
    return rows


def full_campaign_summary() -> dict | None:
    report_path = BENCH / "runs" / PILOT_CAMPAIGN_ID / f"{PILOT_CAMPAIGN_ID}.json"
    report = load_json(report_path)
    if not isinstance(report, dict):
        return None
    summary = report.get("summary") or {}
    return {
        "campaign_id": report.get("campaign_id", PILOT_CAMPAIGN_ID),
        "weighted": summary.get("weighted_composite_by_model") or {},
        "by_model_task": summary.get("by_model_task_composite_median") or {},
        "fractures": summary.get("fracture_counts") or {},
        "runs_scored": summary.get("scored", 0),
        "headline_tasks": summary.get("headline_tasks") or [],
    }


def load_leaderboard_data() -> dict | None:
    """Build leaderboard from scored campaign report (falls back to checked-in JSON)."""
    campaign_path = BENCH / "campaigns" / f"{PILOT_CAMPAIGN_ID}.json"
    report_path = BENCH / "runs" / PILOT_CAMPAIGN_ID / f"{PILOT_CAMPAIGN_ID}.json"
    campaign = load_json(campaign_path)
    report = load_json(report_path)
    if isinstance(campaign, dict) and isinstance(report, dict) and report.get("runs"):
        scripts_dir = str(BENCH / "scripts")
        if scripts_dir not in sys.path:
            sys.path.insert(0, scripts_dir)
        from generate_leaderboard import build_leaderboard  # noqa: WPS433

        return build_leaderboard(campaign, report)
    lb = load_json(BENCH / "docs" / "LEADERBOARD_v0.json")
    return lb if isinstance(lb, dict) else None


def task_label(task_id: str) -> str:
    return TASK_SHORT_NAMES.get(task_id, task_id.replace("_", " "))


def load_campaign_report() -> dict | None:
    report = load_json(BENCH / "runs" / PILOT_CAMPAIGN_ID / f"{PILOT_CAMPAIGN_ID}.json")
    return report if isinstance(report, dict) else None


def model_rank_order(leaderboard: dict | None) -> list[str]:
    if leaderboard and leaderboard.get("rankings"):
        return [row["model_id"] for row in leaderboard["rankings"]]
    return []


def summarize_issue(
    runs: list[dict],
    fracture_labels: dict[str, str],
) -> str:
    composites = [r.get("composite_score") for r in runs if r.get("composite_score") is not None]
    if composites and min(composites) >= 0.999:
        codes: Counter[str] = Counter()
        for rec in runs:
            for code in rec.get("fracture_codes") or []:
                codes[code] += 1
            comp = rec.get("composite") or {}
            for layer in ("l1", "l2", "l3"):
                for code in comp.get(layer, {}).get("fracture_codes") or []:
                    codes[code] += 1
        if not codes:
            return "None"

    codes = Counter()
    failure_modes: Counter[str] = Counter()
    for rec in runs:
        for code in rec.get("fracture_codes") or []:
            codes[code] += 1
        for fm in rec.get("failure_modes") or []:
            failure_modes[fm] += 1
        comp = rec.get("composite") or {}
        for layer in ("l1", "l2", "l3"):
            for code in comp.get(layer, {}).get("fracture_codes") or []:
                codes[code] += 1
            for fm in comp.get(layer, {}).get("failure_modes") or []:
                failure_modes[fm] += 1

    if not codes and not failure_modes:
        if composites and median(composites) >= 0.95:
            return "Minor citation variance"
        return "Score below leader; no dominant fracture code"

    parts: list[str] = []
    for code, count in codes.most_common(2):
        label = fracture_labels.get(code, code.replace("_", " ").lower())
        parts.append(f"{label} ({code})")
    if failure_modes and not codes:
        parts.append(failure_modes.most_common(1)[0][0].replace("_", " "))
    issue = "; ".join(parts)
    if len(codes) > 2:
        issue += f"; +{len(codes) - 2} more"
    return issue


def findings_table_rows(
    report: dict,
    leaderboard: dict | None,
) -> list[dict]:
    fracture_labels = (leaderboard or {}).get("fracture_labels") or {}
    from task_registry import headline_task_ids  # noqa: WPS433

    task_order = report.get("tasks") or []
    model_order = model_rank_order(leaderboard) or report.get("models") or []
    headline_tasks = set(
        (leaderboard or {}).get("methodology", {}).get("headline_tasks")
        or (report.get("summary") or {}).get("headline_tasks")
        or headline_task_ids(task_order)
    )

    headline_by_model: dict[str, float] = {}
    if leaderboard and leaderboard.get("rankings"):
        for row in leaderboard["rankings"]:
            if row.get("headline_composite") is not None:
                headline_by_model[row["model_id"]] = float(row["headline_composite"])
    else:
        for model, score in ((report.get("summary") or {}).get("weighted_composite_by_model") or {}).items():
            headline_by_model[model] = float(score)

    grouped: dict[tuple[str, str], list[dict]] = {}
    for rec in report.get("runs", []):
        if rec.get("status") != "scored":
            continue
        key = (rec["task_id"], rec["model_id"])
        grouped.setdefault(key, []).append(rec)

    rows: list[dict] = []
    for tid in task_order:
        for mid in model_order:
            recs = grouped.get((tid, mid))
            if not recs:
                continue
            l1_scores = [r["composite"]["l1"]["score"] for r in recs if r.get("composite")]
            composites = [r["composite_score"] for r in recs if r.get("composite_score") is not None]
            rows.append({
                "task_id": tid,
                "model_id": mid,
                "l1_median": median(l1_scores) if l1_scores else None,
                "composite_median": median(composites) if composites else None,
                "headline_composite": headline_by_model.get(mid),
                "issue": summarize_issue(recs, fracture_labels),
                "in_headline": tid in headline_tasks,
            })
    return rows


def truncate_cell(text: str, limit: int = 140) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def add_findings_section(
    doc: Document,
    eval_full: dict | None,
    leaderboard: dict | None,
    report: dict | None,
) -> None:
    if not eval_full or not eval_full.get("runs_scored") or not report:
        add_para(doc, "Full 5-task eval not yet scored.")
        return

    add_para(
        doc,
        f"Campaign {eval_full['campaign_id']}: {eval_full['runs_scored']} runs scored "
        "(5 tasks x 3 models x 3 runs). Score = median L1 hard-accuracy; "
        "Composite = median full L1+L2+L3 (3 runs per cell).",
    )

    rows = findings_table_rows(report, leaderboard)
    if not rows:
        add_para(doc, "No scored runs in campaign report.")
        return

    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    headers = ("Task", "Model", "Score", "Composite", "Issue")
    for i, label in enumerate(headers):
        table.rows[0].cells[i].text = label

    for i, row in enumerate(rows, start=1):
        task_txt = task_label(row["task_id"])
        if not row["in_headline"]:
            task_txt += " *"
        table.rows[i].cells[0].text = task_txt
        table.rows[i].cells[1].text = row["model_id"]
        table.rows[i].cells[2].text = (
            f"{row['l1_median']:.3f}" if row["l1_median"] is not None else "-"
        )
        table.rows[i].cells[3].text = (
            f"{row['composite_median']:.3f}" if row["composite_median"] is not None else "-"
        )
        table.rows[i].cells[4].text = truncate_cell(row["issue"])

    add_para(doc, "* Not in headline rank (see note below).", bold=False)

    add_para(
        doc,
        "Why GOOGL is scored but excluded from headline rank: GOOGL_footnote_reconciliation "
        "is included in the full 5-task campaign as a ceiling check — all three models median "
        "1.0 on L1, path recall, and composite across three runs, so it no longer separates "
        "frontier performers. Headline rank therefore uses PEP, AMZN, NFLX, and KO, where "
        "citation quality (CITE_BROAD, CITE_HALLUC) and task-specific traps still produce "
        "meaningful spread (Claude 0.994 vs GPT-4o 0.896 in this pilot).",
    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_pilot_brief() -> Document:
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    commit = git_head()
    tasks = task_summaries()
    eval_full = full_campaign_summary()
    leaderboard = load_leaderboard_data()
    report = load_campaign_report()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    title = doc.add_heading("Zstate Pilot Status Brief", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Work in progress — {today} · git {commit}").italic = True

    add_para(
        doc,
        "Internal team share. Covers architecture, pilot tasks, reward design, PM simulator, "
        "and early model findings. Not for external publication.",
    )

    add_heading(doc, "1. What we are building", 1)
    add_para(
        doc,
        "Zstate is an equity-research benchmark with two complementary tracks in one repo: "
        "Track A (public leaderboard, single-turn forensics/modeling tasks) and Track B "
        "(RL training env with a skeptical Portfolio Manager). Both share fixed EDGAR corpora, "
        "expert ground truth, and a fracture taxonomy.",
    )

    add_heading(doc, "Architecture (high level)", 2)
    diagram = (
        "Shared foundation: Corpus (fixed doc bundles) + Expert pipeline (CFA draft/review) "
        "+ Trajectory schema v1\n"
        "    |\n"
        "    +-- Track A (benchmark_v0.1/) -- single-turn Type F/M tasks\n"
        "    |       -> 3-layer reward (L1 accuracy / L2 gold path / L3 citations)\n"
        "    |       -> Leaderboard + fracture codes\n"
        "    |\n"
        "    +-- Track B (env_v1/) -- dual-control episodes\n"
        "            -> PM policy FSM (pushback, follow-ups, pushover trap)\n"
        "            -> 4-component reward (Outcome / Grounding / Defense / Hallucination)\n"
        "    |\n"
        "    +-- Track C (future export/) -- curated JSONL + reward vectors for training product"
    )
    add_para(doc, diagram)

    add_heading(doc, "2. Reward systems", 1)

    add_heading(doc, "Track A — 3-layer benchmark score", 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = ("Layer", "What it measures", "Typical weight (Type F)")
    for i, label in enumerate(hdr):
        table.rows[0].cells[i].text = label
    rows = [
        ("L1 Hard accuracy", "Python verify vs expert ground truth (metrics, arithmetic)", "55%"),
        ("L2 Gold path", "Section recall, tool order, workflow fidelity", "25%"),
        ("L3 Trust / citations", "Verbatim snippets, policy acks, anti-hallucination", "20%"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    add_para(doc, "Composite = weighted sum. Fracture codes tag failure modes (e.g. CITE_BROAD).")

    add_heading(doc, "Track B — 4-component env reward", 2)
    add_para(
        doc,
        "Reward = 0.45·Outcome + 0.25·Grounding + 0.20·Defense − 0.10·Hallucination. "
        "Defense is env-only: rewards substantive engagement when the PM pushes back. "
        "Capitulation without retrieval triggers pushover penalties.",
    )

    add_heading(doc, "PM simulator (Track B)", 2)
    add_bullets(doc, [
        "Scripted FSM (pm_v1_1): opening skepticism -> follow-up A/B/C branches -> pushover trap.",
        "Agent tools: fixed corpus search + send_message_to_pm + submit_recommendation.",
        "Sample episode: Solaris earnings-quality dispute (env_v1/episodes/).",
        "Defense is the new discriminator vs single-turn QA — see frontier runs in env_v1/runs/.",
    ])

    add_heading(doc, "3. Published pilot tasks (5/15 MVD)", 1)
    ttable = doc.add_table(rows=len(tasks) + 1, cols=5)
    ttable.style = "Table Grid"
    for i, label in enumerate(("Task ID", "Ticker", "Type", "Tier", "Objective")):
        ttable.rows[0].cells[i].text = label
    for i, t in enumerate(tasks, start=1):
        ttable.rows[i].cells[0].text = t["task_id"]
        ttable.rows[i].cells[1].text = t["ticker"]
        ttable.rows[i].cells[2].text = t["type"]
        ttable.rows[i].cells[3].text = t["tier"]
        summary = t["summary"]
        if len(summary) > 120:
            summary = summary[:117] + "..."
        ttable.rows[i].cells[4].text = summary

    add_heading(doc, "Task highlights", 2)
    add_bullets(doc, [
        "GOOGL (F): Q1 2026 segment sum + hedging reconciling item — ceiling task for frontier models.",
        "PEP (M): FY2025 FX / organic growth decomposition from MD&A tables.",
        "AMZN (F): FY2025 segment operating income bridge with stock-based comp decoy.",
        "NFLX (F): FY2025 cash content guidance vs nine-month YTD actuals (guidance drift).",
        "KO (F, new): FY2025 five-segment Total net revenues + Corporate + Eliminations bridge "
        "to $47,941M consolidated; LatAm (2)% total vs (12)% FX from MD&A.",
    ])

    add_heading(doc, "4. Eval methodology", 1)
    add_bullets(doc, [
        "Campaign runner: benchmark_v0.1/scripts/run_benchmark_campaign.py",
        "Eval mode ON: generic citation rules only (no task-specific cheat-sheets).",
        "Path roles: canonical slugs (segment_financials, narrative_fx) — not issuer note numbers.",
        "Headline composite excludes GOOGL (ceiling); PEP + AMZN + NFLX (+ KO in 5-task campaigns).",
        "Models routed: OpenAI (gpt-*), Anthropic (claude-*), Gemini (gemini-* via OpenAI-compatible API).",
        "3 runs per task (median) for all models in pilot_eval_5task_v1.",
    ])

    add_heading(doc, "5. Findings (July 2026 pilot)", 1)
    add_findings_section(doc, eval_full, leaderboard, report)

    add_heading(doc, "6. WIP / next steps", 1)
    add_bullets(doc, [
        "Scale to 15-task MVD (5 companies x 3 task types) with expert sign-off pipeline.",
        "Tighten L3 citation rules or eval-mode prompts — CITE_BROAD is the main separator on headline tasks.",
        "Track B: more frontier trajectories + PM branch coverage.",
        "Expert Workbench (spec) for CFA draft -> review -> publish workflow.",
        "LATER-06: full corpus ingest with excerpt SHA pins (P3-10 mitigation on KO bundle).",
    ])

    add_para(doc, "References: docs/ARCHITECTURE.md · benchmark_v0.1/docs/PILOT_EVAL_JUL2026.md · env_v1/docs/METHODOLOGY_RL_ENV.md")
    return doc


class PDFReport(FPDF):
    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def sanitize_for_pdf(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2192": "->",
        "\u00b7": " ",
        "\u2265": ">=",
        "\u2264": "<=",
        "\u2212": "-",
        "\u00d7": "x",
        "\u00b1": "+/-",
        "\u2022": "-",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("ascii", "replace").decode("ascii")


def pdf_multiline(pdf: FPDF, height: int, text: str) -> None:
    if not text:
        return
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.epw, height, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def docx_to_pdf(doc_path: Path, pdf_path: Path, title: str) -> None:
    from docx import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def iter_block_items(document: DocxDocument):
        parent = document.element.body
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    d = DocxDocument(str(doc_path))
    pdf = PDFReport()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf_multiline(pdf, 10, sanitize_for_pdf(title))
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)

    for block in iter_block_items(d):
        if isinstance(block, Paragraph):
            text = sanitize_for_pdf(block.text.strip())
            if not text:
                pdf.ln(3)
                continue
            style = block.style.name if block.style else ""
            if "Heading 1" in style:
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 14)
                pdf_multiline(pdf, 8, text)
                pdf.set_font("Helvetica", "", 10)
            elif "Heading 2" in style:
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 12)
                pdf_multiline(pdf, 7, text)
                pdf.set_font("Helvetica", "", 10)
            elif "Heading 3" in style:
                pdf.set_font("Helvetica", "B", 11)
                pdf_multiline(pdf, 6, text)
                pdf.set_font("Helvetica", "", 10)
            else:
                pdf_multiline(pdf, 5, text)
        else:
            pdf.ln(2)
            for row in block.rows:
                line = sanitize_for_pdf(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
                if len(line) > 500:
                    line = line[:497] + "..."
                pdf.set_font("Helvetica", "", 8)
                pdf_multiline(pdf, 4, line)
            pdf.set_font("Helvetica", "", 10)
            pdf.ln(2)

    pdf.output(str(pdf_path))


def main() -> None:
    EXPORT.mkdir(parents=True, exist_ok=True)
    doc = build_pilot_brief()
    docx_path = EXPORT / "ZSTATE_Pilot_Status_Brief_Jul2026.docx"
    pdf_path = EXPORT / "ZSTATE_Pilot_Status_Brief_Jul2026.pdf"
    doc.save(docx_path)
    docx_to_pdf(docx_path, pdf_path, "Zstate Pilot Status Brief — July 2026 (WIP)")
    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")


if __name__ == "__main__":
    main()
